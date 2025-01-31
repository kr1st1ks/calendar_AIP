from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL


def export_schedule_to_docx(schedule, file_name):
    doc = Document()

    # Убираем отступы и маргиналы в документе
    section = doc.sections[0]
    section.left_margin = Inches(0.5)  # Убираем отступ слева
    section.right_margin = Inches(0.5)  # Устанавливаем отступ справа в 0.5 дюйма
    section.top_margin = Inches(0.5)  # Убираем верхний отступ
    section.bottom_margin = Inches(0.5)  # Убираем нижний отступ

    # Добавляем заголовок
    doc.add_heading("Расписание", level=0)

    # Создаем таблицу с заголовками
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'  # Устанавливаем стиль таблицы (с рамками)
    table.autofit = False  # Отключаем авторазмеры колонок

    # Настраиваем заголовки таблицы
    headers = table.rows[0].cells
    headers[0].text = "Дата"
    headers[1].text = "Время начала"
    headers[2].text = "Время окончания"
    headers[3].text = "Тема"
    headers[4].text = "Описание"

    # Применяем стили к заголовкам
    for cell in headers:
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.runs[0]
        run.bold = True
        run.font.size = Pt(11)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Добавляем данные в таблицу
    for date, events in sorted(schedule.items()):
        for event in events:
            row_cells = table.add_row().cells
            row_cells[0].text = date
            row_cells[1].text = event.get('start_time', "")
            row_cells[2].text = event.get('end_time', "")
            row_cells[3].text = event.get('theme', "")
            row_cells[4].text = event.get('description', "")

            # Применяем стили к ячейкам
            for cell in row_cells:
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.runs[0]
                run.font.size = Pt(10)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Сохраняем документ
    doc.save(file_name)
