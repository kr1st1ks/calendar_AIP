import sys
import json
import os
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QVBoxLayout, QHBoxLayout, QWidget,
    QPushButton, QLabel, QLineEdit, QCalendarWidget, QTableWidget,
    QTableWidgetItem, QTimeEdit, QTextEdit, QMessageBox, QComboBox,
    QFileDialog
)
from PyQt5.QtCore import QTime, QDate

from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL

from collections import defaultdict


class ScheduleApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Планировщик расписания")
        self.setGeometry(100, 100, 800, 600)

        # Словарь для хранения расписания
        self.schedule = defaultdict(list)

        # Основной виджет
        self.main_widget = QWidget()
        self.setCentralWidget(self.main_widget)
        self.layout = QVBoxLayout()
        self.main_widget.setLayout(self.layout)

        # Верхняя панель с кнопками
        self.button_layout = QHBoxLayout()
        self.layout.addLayout(self.button_layout)

        self.add_event_button = QPushButton("Добавить событие")
        self.add_event_button.setStyleSheet(
            "background-color: #4CAF50; color: white; font-size: 14px; padding: 10px; border-radius: 5px;")
        self.add_event_button.clicked.connect(self.show_add_event_dialog)
        self.button_layout.addWidget(self.add_event_button)

        self.view_schedule_button = QPushButton("Просмотреть расписание")
        self.view_schedule_button.setStyleSheet(
            "background-color: #2196F3; color: white; font-size: 14px; padding: 10px; border-radius: 5px;")
        self.view_schedule_button.clicked.connect(self.show_view_schedule_dialog)
        self.button_layout.addWidget(self.view_schedule_button)

        self.export_button = QPushButton("Экспортировать в .docx")
        self.export_button.setStyleSheet(
            "background-color: #FFC107; color: white; font-size: 14px; padding: 10px; border-radius: 5px;")
        self.export_button.clicked.connect(self.export_to_docx)
        self.button_layout.addWidget(self.export_button)

        # Календарь
        self.calendar = QCalendarWidget()
        self.calendar.setStyleSheet("""
            QCalendarWidget {
                background-color: #ffffff;
                border: 1px solid #ddd;
                font-size: 14px;
            }
            QCalendarWidget QAbstractItemView {
                selection-background-color: #4CAF50;
            }
            QCalendarWidget QToolButton {
                background-color: #4CAF50;
                color: white;
                font-weight: bold;
                border: none;
                padding: 10px;
            }
            QCalendarWidget QToolButton:hover {
                background-color: #45a049;
            }
            QCalendarWidget QToolButton:pressed {
                background-color: #3e8e41;
            }
            QCalendarWidget QHeaderView::section {
                background-color: #2196F3;
                color: white;
                font-size: 16px;
                padding: 10px;
                border: none;
                text-align: center;
            }
        """)
        self.calendar.clicked.connect(self.update_schedule_view)
        self.layout.addWidget(self.calendar)

        # Таблица для отображения событий
        self.event_table = QTableWidget()
        self.event_table.setColumnCount(5)  # Увеличиваем количество столбцов
        self.event_table.setHorizontalHeaderLabels(["Дата", "Начало", "Конец", "Тема", "Описание"])  # Добавляем "Тема"
        self.event_table.setStyleSheet("""
            QTableWidget {
                border: 1px solid #ddd;
                font-size: 14px;
                background-color: #f9f9f9;
                gridline-color: #ddd;
            }
            QTableWidget::item {
                padding: 10px;
            }
            QTableWidget::item:hover {
                background-color: #f1f1f1;
            }
            QHeaderView::section {
                background-color: #4CAF50;
                color: white;
                font-size: 16px;
                padding: 10px;
            }
        """)
        self.layout.addWidget(self.event_table)

        # Кнопка для удаления события
        self.delete_event_button = QPushButton("Удалить событие")
        self.delete_event_button.setStyleSheet(
            "background-color: #f44336; color: white; font-size: 14px; padding: 10px; border-radius: 5px;")
        self.delete_event_button.clicked.connect(self.delete_event)
        self.layout.addWidget(self.delete_event_button)

        # Кнопка для редактирования события
        self.edit_event_button = QPushButton("Редактировать событие")
        self.edit_event_button.setStyleSheet(
            "background-color: #009688; color: white; font-size: 14px; padding: 10px; border-radius: 5px;")
        self.edit_event_button.clicked.connect(self.edit_event)
        self.layout.addWidget(self.edit_event_button)

        self.search_layout = QHBoxLayout()
        self.layout.addLayout(self.search_layout)

        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Введите текст для поиска...")
        self.search_layout.addWidget(self.search_input)

        self.search_button = QPushButton("Поиск")
        self.search_button.setStyleSheet(
            "background-color: #FFA726; color: white; font-size: 14px; padding: 10px; border-radius: 5px;")
        self.search_button.clicked.connect(self.search_events)
        self.search_layout.addWidget(self.search_button)
        self.search_input.returnPressed.connect(self.search_button.click)

        # Обновить расписание для выбранной даты
        self.update_schedule_view()

        self.load_schedule_from_file()
        self.update_schedule_view()

    def load_schedule_from_file(self):
        if os.path.exists("schedule.json"):
            try:
                with open("schedule.json", "r", encoding="utf-8") as file:
                    # Загружаем расписание из файла
                    data = json.load(file)
                    # Конвертируем обратно в defaultdict
                    self.schedule = defaultdict(list, data)
                print("Расписание загружено из файла.")
            except Exception as e:
                QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить расписание: {e}")

    def show_add_event_dialog(self):
        # Диалог для добавления события
        dialog = QWidget()
        dialog.setWindowTitle("Добавить событие")
        dialog_layout = QVBoxLayout()
        dialog.setLayout(dialog_layout)

        # Ввод даты
        date_label = QLabel("Дата:")
        dialog_layout.addWidget(date_label)
        date_input = QLineEdit(self.calendar.selectedDate().toString("yyyy-MM-dd"))
        dialog_layout.addWidget(date_input)

        # Ввод времени начала
        start_time_label = QLabel("Время начала:")
        dialog_layout.addWidget(start_time_label)
        start_time_input = QTimeEdit()
        start_time_input.setTime(QTime(9, 0))
        dialog_layout.addWidget(start_time_input)

        # Ввод времени окончания
        end_time_label = QLabel("Время окончания:")
        dialog_layout.addWidget(end_time_label)
        end_time_input = QTimeEdit()
        end_time_input.setTime(QTime(10, 0))
        dialog_layout.addWidget(end_time_input)

        # Ввод описания
        description_label = QLabel("Описание:")
        dialog_layout.addWidget(description_label)
        description_input = QTextEdit()
        dialog_layout.addWidget(description_input)

        # Добавьте виджет выбора темы
        theme_label = QLabel("Тема:")
        dialog_layout.addWidget(theme_label)

        # Создаем комбинированный виджет для выбора темы или ввода новой
        theme_input = QComboBox()
        theme_input.setEditable(True)  # Позволяем редактировать текст
        theme_input.addItem("")  # Строка-заполнитель

        # Собираем все уникальные темы из расписания
        all_themes = set()
        for date, events in self.schedule.items():
            for event in events:
                all_themes.add(event['theme'])

        # Добавляем существующие темы в комбинированный список
        for theme in sorted(all_themes):  # Сортируем для удобства
            theme_input.addItem(theme)

        dialog_layout.addWidget(theme_input)

        # Кнопки "Добавить" и "Отмена"
        button_layout = QHBoxLayout()
        dialog_layout.addLayout(button_layout)
        add_button = QPushButton("Добавить")
        cancel_button = QPushButton("Отмена")
        button_layout.addWidget(add_button)
        button_layout.addWidget(cancel_button)

        def add_event():
            date = date_input.text()
            start_time = start_time_input.time().toString("HH:mm")
            end_time = end_time_input.time().toString("HH:mm")
            theme = theme_input.currentText().strip()  # Получаем тему
            description = description_input.toPlainText()

            # Проверка накладок
            for event in self.schedule[date]:
                if not (end_time <= event['start_time'] or start_time >= event['end_time']):
                    reply = QMessageBox.question(
                        dialog,
                        "Накладка",
                        f"Событие пересекается с '{event['description']}' "
                        f"({event['start_time']} - {event['end_time']}). Добавить всё равно?",
                        QMessageBox.Yes | QMessageBox.No
                    )
                    if reply == QMessageBox.No:
                        return

            # Добавление события
            self.schedule[date].append({
                'start_time': start_time,
                'end_time': end_time,
                'theme': theme,  # Сохраняем тему
                'description': description
            })
            QMessageBox.information(dialog, "Успех", "Событие добавлено!")
            dialog.close()
            self.update_schedule_view()

        add_button.clicked.connect(add_event)
        cancel_button.clicked.connect(dialog.close)
        dialog.show()

    def delete_event(self):
        # Получаем выбранную строку в таблице
        selected_row = self.event_table.currentRow()
        if selected_row == -1:  # Если строка не выбрана
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите событие для удаления.")
            return

        # Получаем данные из выбранной строки
        start_time = self.event_table.item(selected_row, 1).text()
        end_time = self.event_table.item(selected_row, 2).text()
        theme = self.event_table.item(selected_row, 3).text()
        description = self.event_table.item(selected_row, 4).text()

        # Получаем выбранную дату из календаря
        selected_date = self.calendar.selectedDate().toString("yyyy-MM-dd")

        # Ищем и удаляем событие из словаря
        for event in self.schedule[selected_date]:
            if (event['start_time'] == start_time and
                    event['end_time'] == end_time and
                    event['description'] == description):
                self.schedule[selected_date].remove(event)
                QMessageBox.information(self, "Удаление", "Событие успешно удалено.")
                break
        else:
            QMessageBox.warning(self, "Ошибка", "Не удалось найти событие для удаления.")

        # Обновление таблицы
        self.update_schedule_view()

    def edit_event(self):
        # Получаем выбранную строку в таблице
        selected_row = self.event_table.currentRow()
        if selected_row == -1:  # Если строка не выбрана
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите событие для редактирования.")
            return

        # Получаем данные из выбранной строки
        start_time = self.event_table.item(selected_row, 1).text()
        end_time = self.event_table.item(selected_row, 2).text()
        theme = self.event_table.item(selected_row, 3).text()
        description = self.event_table.item(selected_row, 4).text()

        # Получаем выбранную дату из календаря
        selected_date = self.calendar.selectedDate().toString("yyyy-MM-dd")

        # Открываем диалог редактирования
        dialog = QWidget()
        dialog.setWindowTitle("Редактировать событие")
        dialog_layout = QVBoxLayout()
        dialog.setLayout(dialog_layout)

        # Ввод времени начала
        start_time_label = QLabel("Время начала:")
        dialog_layout.addWidget(start_time_label)
        start_time_input = QTimeEdit(QTime.fromString(start_time, "HH:mm"))
        dialog_layout.addWidget(start_time_input)

        # Ввод времени окончания
        end_time_label = QLabel("Время окончания:")
        dialog_layout.addWidget(end_time_label)
        end_time_input = QTimeEdit(QTime.fromString(end_time, "HH:mm"))
        dialog_layout.addWidget(end_time_input)

        # Ввод описания
        description_label = QLabel("Описание:")
        dialog_layout.addWidget(description_label)
        description_input = QTextEdit(description)
        dialog_layout.addWidget(description_input)

        # Ввод темы
        theme_label = QLabel("Тема:")
        dialog_layout.addWidget(theme_label)

        # Создаем комбинированный виджет для выбора темы или ввода новой
        theme_input = QComboBox()
        theme_input.setEditable(True)  # Позволяем редактировать текст
        theme_input.addItem("Выберите или введите тему...")  # Строка-заполнитель

        # Собираем все уникальные темы из расписания
        all_themes = set()
        for date, events in self.schedule.items():
            for event in events:
                all_themes.add(event['theme'])

        # Добавляем существующие темы в комбинированный список
        for theme in sorted(all_themes):  # Сортируем для удобства
            theme_input.addItem(theme)

        theme_input.setCurrentText(theme)  # Устанавливаем текущую тему
        dialog_layout.addWidget(theme_input)

        # Кнопки "Сохранить" и "Отмена"
        button_layout = QHBoxLayout()
        dialog_layout.addLayout(button_layout)
        save_button = QPushButton("Сохранить")
        cancel_button = QPushButton("Отмена")
        button_layout.addWidget(save_button)
        button_layout.addWidget(cancel_button)

        def save_changes():
            new_start_time = start_time_input.time().toString("HH:mm")
            new_end_time = end_time_input.time().toString("HH:mm")
            new_description = description_input.toPlainText()
            new_theme = theme_input.currentText().strip()  # Получаем новую тему
            # Проверка накладок с другими событиями
            # for event in self.schedule[selected_date]:
            #    if not (new_end_time <= event['start_time'] or new_start_time >= event['end_time']):
            #        reply = QMessageBox.question(
            #            dialog,
            #            "Накладка",
            #            f"Событие пересекается с '{event['description']}' "
            #            f"({event['start_time']} - {event['end_time']}). Сохранить изменения?",
            #            QMessageBox.Yes | QMessageBox.No
            #       )
            #        if reply == QMessageBox.No:
            #            return
            # Проверка накладок с другими событиями
            for event in self.schedule[selected_date]:
                if not (new_end_time <= event['start_time'] or new_start_time >= event['end_time']):
                    reply = QMessageBox.question(
                        dialog,
                        "Накладка",
                        f"Событие пересекается с '{event['description']}' "
                        f"({event['start_time']} - {event['end_time']}). Сохранить изменения?",
                        QMessageBox.Yes | QMessageBox.No
                    )
                    if reply == QMessageBox.No:
                        return

            # Обновление события
            for event in self.schedule[selected_date]:
                if (event['start_time'] == start_time and
                        event['end_time'] == end_time and
                        event['description'] == description):
                    event['start_time'] = new_start_time
                    event['end_time'] = new_end_time
                    event['theme'] = new_theme  # Обновляем тему
                    event['description'] = new_description
                    break

            # Сортировка событий по времени начала
            self.schedule[selected_date].sort(key=lambda event: event['start_time'])

            QMessageBox.information(dialog, "Успех", "Событие успешно обновлено.")
            dialog.close()
            self.update_schedule_view()

        save_button.clicked.connect(save_changes)
        cancel_button.clicked.connect(dialog.close)
        dialog.show()

    def show_view_schedule_dialog(self):
        # Проверка, открыто ли уже окно расписания
        if hasattr(self, "view_schedule_dialog") and self.view_schedule_dialog is not None:
            self.view_schedule_dialog.close()

        # Диалог для просмотра расписания
        self.view_schedule_dialog = QWidget()
        self.view_schedule_dialog.setWindowTitle("Просмотреть расписание")
        dialog_layout = QVBoxLayout()
        self.view_schedule_dialog.setLayout(dialog_layout)

        # Добавление выбора начальной даты
        filter_layout = QHBoxLayout()
        dialog_layout.addLayout(filter_layout)

        # Календарь для выбора начальной даты
        start_date_label = QLabel("Выберите дату начала:")
        filter_layout.addWidget(start_date_label)
        start_date_input = QCalendarWidget()
        start_date_input.setSelectedDate(QDate.currentDate())  # По умолчанию текущая дата
        filter_layout.addWidget(start_date_input)

        # Выпадающий список для выбора диапазона
        range_label = QLabel("Выберите диапазон:")
        filter_layout.addWidget(range_label)

        date_range_input = QComboBox()  # Выпадающий список для диапазона
        date_range_input.addItem("1 неделя")  # Список опций
        date_range_input.addItem("1 месяц")
        date_range_input.addItem("3 месяца")
        filter_layout.addWidget(date_range_input)

        # Добавление выбора темы
        theme_filter_label = QLabel("Фильтр по теме:")
        filter_layout.addWidget(theme_filter_label)

        theme_filter_input = QComboBox()  # Выпадающий список для выбора темы
        theme_filter_input.addItem("Все темы")  # По умолчанию показываем все темы

        # Собираем все уникальные темы из расписания
        all_themes = set()
        for date, events in self.schedule.items():
            for event in events:
                all_themes.add(event['theme'])

        for theme in sorted(all_themes):  # Сортируем для удобства
            theme_filter_input.addItem(theme)

        filter_layout.addWidget(theme_filter_input)

        # Кнопка для применения фильтра
        filter_button = QPushButton("Применить фильтр")
        filter_layout.addWidget(filter_button)

        # Таблица для вывода расписания
        schedule_table = QTableWidget()
        schedule_table.setColumnCount(5)
        schedule_table.setHorizontalHeaderLabels(["Дата", "Время начала", "Время окончания", "Тема", "Описание"])
        dialog_layout.addWidget(schedule_table)

        def populate_table(filtered_schedule=None):
            """Обновляет таблицу расписания, фильтруя по выбранной дате и диапазону."""
            schedule_table.setRowCount(0)  # Очистить таблицу

            # Если фильтр пустой, показываем всё расписание
            data = self.schedule.items() if filtered_schedule is None else filtered_schedule.items()

            # Применяем фильтрацию по дате и диапазону
            selected_start_date = start_date_input.selectedDate()
            selected_range = date_range_input.currentText()

            # Преобразуем диапазон в дни
            if selected_range == "1 неделя":
                end_date = selected_start_date.addDays(7)
            elif selected_range == "1 месяц":
                end_date = selected_start_date.addMonths(1)
            elif selected_range == "3 месяца":
                end_date = selected_start_date.addMonths(3)

            # Фильтрация событий по диапазону
            selected_theme = theme_filter_input.currentText()  # Получаем выбранную тему
            for date, events in sorted(data):
                row_date = QDate.fromString(date, "yyyy-MM-dd")
                if selected_start_date <= row_date <= end_date:
                    for event in sorted(events, key=lambda e: e['start_time']):
                        # Применяем фильтрацию по теме
                        if selected_theme == "Все темы" or event['theme'] == selected_theme:
                            row = schedule_table.rowCount()
                            schedule_table.insertRow(row)
                            schedule_table.setItem(row, 0, QTableWidgetItem(date))
                            schedule_table.setItem(row, 1, QTableWidgetItem(event['start_time']))
                            schedule_table.setItem(row, 2, QTableWidgetItem(event['end_time']))
                            schedule_table.setItem(row, 3, QTableWidgetItem(event['theme']))
                            schedule_table.setItem(row, 4, QTableWidgetItem(event['description']))

        def apply_filter():
            """Применяет фильтр по выбранной дате, диапазону и теме."""
            filtered_schedule = {}

            # Получаем выбранные даты для диапазона
            selected_start_date = start_date_input.selectedDate().toString("yyyy-MM-dd")
            selected_range = date_range_input.currentText()

            # Преобразуем диапазон в дни
            if selected_range == "1 неделя":
                end_date = start_date_input.selectedDate().addDays(7)
            elif selected_range == "1 месяц":
                end_date = start_date_input.selectedDate().addMonths(1)
            elif selected_range == "3 месяца":
                end_date = start_date_input.selectedDate().addMonths(3)

            # Фильтрация по дате и диапазону
            for date, events in self.schedule.items():
                row_date = QDate.fromString(date, "yyyy-MM-dd")
                if start_date_input.selectedDate() <= row_date <= end_date:
                    filtered_schedule[date] = events

            populate_table(filtered_schedule)

        filter_button.clicked.connect(apply_filter)

        # Изначально показываем всё расписание
        populate_table()

        # Кнопка закрытия
        close_button = QPushButton("Закрыть")
        close_button.clicked.connect(self.view_schedule_dialog.close)
        dialog_layout.addWidget(close_button)

        self.view_schedule_dialog.show()

    def search_events(self):
        search_term = self.search_input.text().strip().lower()
        if not search_term:
            QMessageBox.warning(self, "Ошибка", "Введите текст для поиска.")
            return

        filtered_schedule = defaultdict(list)
        for date, events in self.schedule.items():
            for event in events:
                if (search_term in event['theme'].lower() or  # Фильтр по теме
                        search_term in event['description'].lower()):  # Фильтр по описанию
                    filtered_schedule[date].append(event)

        if not filtered_schedule:
            QMessageBox.information(self, "Поиск", f"Ничего не найдено по запросу: {search_term}")
            return

        self.update_schedule_view(filtered_schedule)

    def update_schedule_view(self, filtered_schedule=None):
        """
        Обновляет виджет таблицы расписания.
        Если передан filtered_schedule, отображает только отфильтрованные данные.
        """
        # Если передан словарь с фильтром
        if isinstance(filtered_schedule, dict):
            data = filtered_schedule.items()
        else:
            # Получаем расписание для выбранной даты из календаря
            selected_date = self.calendar.selectedDate().toString("yyyy-MM-dd")
            events = self.schedule.get(selected_date, [])
            data = {selected_date: events}.items()

        # Очищаем таблицу
        self.event_table.setRowCount(0)

        # Заполняем таблицу новыми данными
        for date, events in data:
            for event in events:
                row = self.event_table.rowCount()
                self.event_table.insertRow(row)
                self.event_table.setItem(row, 0, QTableWidgetItem(date))  # Дата
                self.event_table.setItem(row, 1, QTableWidgetItem(event['start_time']))  # Начало
                self.event_table.setItem(row, 2, QTableWidgetItem(event['end_time']))  # Конец
                self.event_table.setItem(row, 3, QTableWidgetItem(event['theme']))  # Тема
                self.event_table.setItem(row, 4, QTableWidgetItem(event['description']))  # Описание

        # Обновляем отображение таблицы
        self.event_table.resizeColumnsToContents()

    def export_to_docx(self):
        # Получаем все данные из расписания
        if not self.schedule:
            QMessageBox.warning(self, "Ошибка", "Нет событий для экспорта.")
            return

        # Запрашиваем путь и имя файла для сохранения
        file_name, _ = QFileDialog.getSaveFileName(self, "Сохранить расписание", "", "Документ Word (*.docx)")

        if not file_name:
            return  # Если пользователь отменил, выходим

        # Создаем новый документ Word
        doc = Document()
        doc.add_heading("Расписание", 0)

        # Создаем таблицу с 5 столбцами: Дата, Время начала, Время окончания, Тема, Описание
        table = doc.add_table(rows=1, cols=5)

        # Заголовки таблицы
        headers = table.rows[0].cells
        headers[0].text = "Дата"
        headers[1].text = "Время начала"
        headers[2].text = "Время окончания"
        headers[3].text = "Тема"
        headers[4].text = "Описание"

        # Применяем стиль к заголовкам
        for cell in headers:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = 1  # Центрируем заголовки

        # Устанавливаем ширину столбцов
        table.columns[0].width = Pt(80)  # Дата
        table.columns[1].width = Pt(80)  # Время начала
        table.columns[2].width = Pt(80)  # Время окончания
        table.columns[3].width = Pt(120)  # Тема
        table.columns[4].width = Pt(20)  # Описание (широкий столбец для описания)

        # Добавляем все события в таблицу
        for date, events in sorted(self.schedule.items()):
            for event in events:
                row_cells = table.add_row().cells
                row_cells[0].text = date
                row_cells[1].text = event['start_time']
                row_cells[2].text = event['end_time']
                row_cells[3].text = event['theme']  # Добавляем тему
                row_cells[4].text = event['description']

                # Настроим выравнивание для текста в ячейках
                row_cells[0].paragraphs[0].alignment = 0  # Выравнивание даты по центру
                row_cells[1].paragraphs[0].alignment = 0  # Выравнивание времени начала по центру
                row_cells[2].paragraphs[0].alignment = 0  # Выравнивание времени окончания по центру
                row_cells[3].paragraphs[0].alignment = 0  # Описание выравниваем по левому краю

                # Сделаем текст в столбце "Описание" многострочным (перенос слов)
                for paragraph in row_cells[4].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)  # Устанавливаем размер шрифта для текста в описании

                # Вертикальное выравнивание для всех ячеек
                for cell in row_cells:
                    cell.paragraphs[0].alignment = 3
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Выравнивание по вертикали

        # Сохраняем файл
        try:
            doc.save(file_name)
            QMessageBox.information(self, "Экспорт", f"Все события успешно экспортированы в {file_name}.")
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось сохранить документ: {e}")

    def save_schedule_to_file(self):
        try:
            with open("schedule.json", "w", encoding="utf-8") as file:
                # Преобразуем в стандартный словарь для сохранения
                json.dump(dict(self.schedule), file, ensure_ascii=False, indent=4)
            print("Расписание сохранено в файл.")
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось сохранить расписание: {e}")

    def closeEvent(self, event):
        self.save_schedule_to_file()
        event.accept()


if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = ScheduleApp()
    window.show()
    sys.exit(app.exec_())
